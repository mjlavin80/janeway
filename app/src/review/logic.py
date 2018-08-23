__copyright__ = "Copyright 2017 Birkbeck, University of London"
__author__ = "Martin Paul Eve & Andy Byers"
__license__ = "AGPL v3"
__maintainer__ = "Birkbeck Centre for Technology and Publishing"

from datetime import timedelta
from uuid import uuid4
import os


from django.urls import reverse
from django.utils import timezone
from django.contrib import messages
from django.conf import settings
from docx import Document

from utils import render_template, setting_handler
from core import models as core_models, files
from review import models
from events import logic as event_logic
from submission import models as submission_models


def get_reviewers(article, request=None):
    review_assignments = article.reviewassignment_set.filter(review_round=article.current_review_round_object())
    reviewers = [review.reviewer.pk for review in review_assignments]
    reviewers.append(request.user.pk)

    return core_models.AccountRole.objects.filter(role__slug='reviewer', journal=request.journal).exclude(
        user__pk__in=reviewers)


def get_suggested_reviewers(article, reviewers):
    suggested_reviewers = []
    keywords = [keyword.word for keyword in article.keywords.all()]
    for reviewer_role in reviewers:
        interests = [interest.name for interest in reviewer_role.user.interest.all()]
        for interest in interests:
            if interest in keywords:
                suggested_reviewers.append(reviewer_role)
                break

    return suggested_reviewers


def get_assignment_content(request, article, editor, assignment):
    email_context = {
        'article': article,
        'editor': editor,
        'assignment': assignment,
    }

    return render_template.get_message_content(request, email_context, 'editor_assignment')


def get_reviewer_notification(request, article, editor, review_assignment):
    review_url = "{0}{1}".format(request.journal_base_url, reverse('do_review',
                                                                   kwargs={'assignment_id': review_assignment.id}))

    access_codes = setting_handler.get_setting('general', 'enable_one_click_access', request.journal).value

    if access_codes:
        review_url = "{0}?access_code={1}".format(review_url, review_assignment.access_code)

    email_context = {
        'article': article,
        'editor': editor,
        'review_assignment': review_assignment,
        'review_url': review_url
    }

    return render_template.get_message_content(request, email_context, 'review_assignment')


def get_withdrawl_notification(request, review_assignment):

    email_context = {
        'article': review_assignment.article,
        'review_assignment': review_assignment,
        'editor': request.user,
    }

    return render_template.get_message_content(request, email_context, 'review_withdrawl')


def get_decision_content(request, article, decision, author_review_url):

    email_context = {
        'article': article,
        'decision': decision,
        'review_url': author_review_url,
    }

    template_name = "review_decision_{0}".format(decision)

    return render_template.get_message_content(request, email_context, template_name)


def get_revision_request_content(request, article, revision):

    email_context = {
        'article': article,
        'revision': revision,
    }

    return render_template.get_message_content(request, email_context, 'request_revisions')


def get_reviewer_from_post(request):
    reviewer_id = request.POST.get('reviewer')

    if reviewer_id:
        reviewer = core_models.Account.objects.get(pk=reviewer_id)

        # if this user is not a reviewer, return None to force an error on the form.
        if not reviewer.is_reviewer(request):
            return None

        return reviewer
    else:
        return None


def log_revision_event(text, user, revision_request):
    action = models.RevisionAction.objects.create(
        text=text,
        logged=timezone.now(),
        user=user
    )

    revision_request.actions.add(action)


def get_draft_email_message(request, article):

    email_context = {
        'article': article,
    }

    return render_template.get_message_content(request, email_context, 'draft_message')


def group_files(article, reviews):
    files = list()

    for file in article.manuscript_files.all():
        files.append(file)

    for file in article.data_figure_files.all():
        files.append(file)

    for review in reviews:
        if review.for_author_consumption and review.display_review_file:
            files.append(review.review_file)

    return files


def handle_decision_action(article, draft, request):
    from submission import models as submission_models
    kwargs = {
        'article': article,
        'request': request,
        'decision': draft.decision,
        'user_message_content': request.POST.get('email_message', 'No message found.'),
        'skip': False,
    }

    if draft.decision == 'accept':
        article.accept_article(stage=submission_models.STAGE_EDITOR_COPYEDITING)
        event_logic.Events.raise_event(event_logic.Events.ON_ARTICLE_ACCEPTED, task_object=article, **kwargs)
    elif draft.decision == 'decline':
        article.decline_article()
        event_logic.Events.raise_event(event_logic.Events.ON_ARTICLE_DECLINED, task_object=article, **kwargs)
    elif draft.decision == 'minor_revisions' or draft.decision == 'major_revisions':
        revision = models.RevisionRequest.objects.create(
            article=article,
            editor=draft.section_editor,
            editor_note='',
            type=draft.decision,
            date_due=timezone.now() + timedelta(days=14)
        )
        article.stage = submission_models.STAGE_UNDER_REVISION
        article.save()

        kwargs['revision'] = revision
        event_logic.Events.raise_event(event_logic.Events.ON_REVISIONS_REQUESTED_NOTIFY, **kwargs)


def get_access_code(request):
    if request.GET.get('access_code'):
        access_code = request.GET.get('access_code')
    else:
        access_code = None

    return access_code


def quick_assign(request, article):
    errors = []
    try:
        default_review_form_id = setting_handler.get_setting('general',
                                                             'default_review_form',
                                                             request.journal).processed_value
    except models.ReviewForm.DoesNotExist:
        errors.append('This journal has no default review form.')

    try:
        review_form = models.ReviewForm.objects.get(pk=default_review_form_id)
    except ValueError:
        errors.append('Default review form is not an integer.')

    try:
        default_visibility = setting_handler.get_setting('general',
                                                         'default_review_visibility',
                                                         request.journal).value
        default_due = setting_handler.get_setting('general',
                                                  'default_review_days',
                                                  request.journal).value
    except BaseException:
        errors.append('This journal does not have either default visibilty or default due.')

    user_id = request.POST.get('quick_assign')
    user = core_models.Account.objects.get(pk=user_id)

    if user not in request.journal.users_with_role('reviewer'):
        errors.append('This user is not a reviewer for this journal.')

    if not errors:
        new_assignment = models.ReviewAssignment.objects.create(
            article=article,
            reviewer=user,
            editor=request.user,
            review_round=article.current_review_round_object(),
            form=review_form,
            access_code=uuid4(),
            visibility=default_visibility,
            date_due=timezone.now() + timedelta(days=int(default_due)),
        )

        article.stage = submission_models.STAGE_UNDER_REVIEW
        article.save()

        email_content = get_reviewer_notification(request, article, request.user, new_assignment)

        kwargs = {'user_message_content': email_content,
                  'review_assignment': new_assignment,
                  'request': request,
                  'skip': False,
                  'acknowledgement': False}

        event_logic.Events.raise_event(event_logic.Events.ON_REVIEWER_REQUESTED, **kwargs)
        event_logic.Events.raise_event(event_logic.Events.ON_REVIEWER_REQUESTED_ACKNOWLEDGE, **kwargs)

        return new_assignment

    else:
        for error in errors:
            messages.add_message(request, messages.WARNING, error)


def handle_reviewer_form(request, new_reviewer_form):
    account = new_reviewer_form.save(commit=False)
    account.username = account.email
    from core import models as core_models
    account.country = core_models.Country.objects.filter(code='GB')[0]
    account.institution = 'N/a'
    account.is_active = True
    account.save()
    account.add_account_role('reviewer', request.journal)
    messages.add_message(request, messages.INFO, 'A new account has been created.')


def get_enrollable_users(request):
    account_roles = core_models.AccountRole.objects.filter(journal=request.journal, role__slug='reviewer')
    users_with_role = [assignment.user.pk for assignment in account_roles]
    return core_models.Account.objects.all().order_by('last_name').exclude(pk__in=users_with_role)


def generate_access_code_url(url_name, assignment, access_code):

    reverse_url = reverse(url_name, kwargs={'assignment_id': assignment.pk})

    if access_code:
        reverse_url = '{reverse_url}?access_code={access_code}'.format(reverse_url=reverse_url, access_code=access_code)

    return reverse_url


def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]


def serve_review_file(assignment):
    """
    Produces a word document representing the review form.
    :param assignment: ReviewAssignment object
    :return: HttpStreamingResponse
    """
    elements = assignment.form.elements.all()
    document = Document()
    document.add_heading('Review #{pk}'.format(pk=assignment.pk), 0)
    document.add_heading('Review of `{article_title}` by {reviewer}'.format(article_title=assignment.article.title,
                                                                            reviewer=assignment.reviewer.full_name()),
                         level=1)
    document.add_paragraph()
    document.add_paragraph('Complete the form below, then upload it under the "FILE UPLOAD" section on your review page'
                           '. There is no need to complete the form on the web page if you are uploading this '
                           'document.')
    document.add_paragraph()

    for element in elements:
        document.add_heading(element.name, level=2)
        document.add_paragraph(element.help_text)
        if element.choices:
            choices = render_choices(element.choices)
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Choice'
            hdr_cells[1].text = 'Indication'

            for choice in element.choices.split('|'):
                row_cells = table.add_row().cells
                row_cells[0].text = str(choice)
        document.add_paragraph()

    filename = '{uuid}.docx'.format(uuid=uuid4())
    filepath = os.path.join(settings.BASE_DIR, 'files', 'temp', filename)
    document.save(filepath)
    return files.serve_temp_file(filepath, filename)


def handle_review_file_switch(review, switch):
    if switch == 'true':
        review.display_review_file = True
    else:
        review.display_review_file = False

    review.save()
